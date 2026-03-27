"""
Export module — Generate Word document from multi-agent advising session.
Week 3, Lonely Octopus AI Agent Bootcamp
"""

from __future__ import annotations

import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


MTM_TEAL = RGBColor(0x0E, 0x74, 0x90)
MTM_DARK = RGBColor(0x1C, 0x48, 0x7B)
MTM_LIGHT = RGBColor(0x85, 0xAB, 0xBD)


def _add_heading(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = MTM_TEAL


def _parse_markdown_to_runs(paragraph, text: str):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            italic_parts = re.split(r"(\*.*?\*)", part)
            for ip in italic_parts:
                if ip.startswith("*") and ip.endswith("*") and not ip.startswith("**"):
                    run = paragraph.add_run(ip[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(ip)


def generate_docx(messages: list[dict], org_profile: dict, agent_logs: dict = None) -> bytes:
    """Generate a Word document from the advising session."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = MTM_DARK

    # Title
    title = doc.add_heading("Technology Advice", level=0)
    for run in title.runs:
        run.font.color.rgb = MTM_TEAL

    org_name = org_profile.get("org_name", "Your Organization")
    subtitle = doc.add_paragraph()
    run = subtitle.add_run(f"Prepared for {org_name}")
    run.font.size = Pt(14)
    run.font.color.rgb = MTM_DARK

    date_para = doc.add_paragraph()
    run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.size = Pt(11)
    run.font.color.rgb = MTM_LIGHT
    run.italic = True

    # Architecture note
    arch_para = doc.add_paragraph()
    run = arch_para.add_run("Multi-Agent Advisory System — Week 3, AI Agent Bootcamp")
    run.font.size = Pt(10)
    run.font.color.rgb = MTM_LIGHT
    run.italic = True

    doc.add_paragraph()

    # Org Profile
    _add_heading(doc, "Organization Profile", level=1)
    profile_labels = {
        "org_name": "Organization",
        "budget_tier": "Annual Budget",
        "staff_count": "Staff Count",
        "cause_area": "Cause Area",
        "current_tech": "Current Technology",
        "pain_points": "Pain Points",
        "it_capacity": "IT Capacity",
    }
    for key, label in profile_labels.items():
        value = org_profile.get(key, "")
        if value:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            p.add_run(str(value))

    doc.add_paragraph()

    # Conversation
    _add_heading(doc, "Advising Session", level=1)

    question_num = 0
    for i, msg in enumerate(messages):
        if msg["role"] == "user":
            question_num += 1
            _add_heading(doc, f"Question {question_num}", level=2)
            doc.add_paragraph(msg["content"])

        elif msg["role"] == "assistant":
            _add_heading(doc, "Advisor Response", level=3)

            # Show which agents were involved
            if agent_logs and i in agent_logs:
                log = agent_logs[i]
                if log.get("routing"):
                    routing = log["routing"]
                    agents_text = f"Specialists consulted: {routing['primary']}"
                    if routing.get("secondary"):
                        agents_text += f", {routing['secondary']}"
                    p = doc.add_paragraph()
                    run = p.add_run(agents_text)
                    run.font.size = Pt(10)
                    run.font.color.rgb = MTM_LIGHT
                    run.italic = True

            content = msg["content"]
            for line in content.split("\n"):
                stripped = line.strip()
                if not stripped:
                    continue

                if stripped.startswith("### "):
                    _add_heading(doc, stripped[4:], level=3)
                elif stripped.startswith("## "):
                    _add_heading(doc, stripped[3:], level=2)
                elif stripped.startswith("# "):
                    _add_heading(doc, stripped[2:], level=1)
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    p = doc.add_paragraph(style="List Bullet")
                    _parse_markdown_to_runs(p, stripped[2:])
                elif re.match(r"^\d+\.\s", stripped):
                    p = doc.add_paragraph(style="List Number")
                    text = re.sub(r"^\d+\.\s", "", stripped)
                    _parse_markdown_to_runs(p, text)
                else:
                    p = doc.add_paragraph()
                    _parse_markdown_to_runs(p, stripped)

            doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    doc.add_paragraph("_" * 60)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "This advice was generated by a multi-agent AI advisory team from Meet the Moment (mtm.now).\n"
        "Specialists: Security Advisor, Technology Advisor, AI Readiness Advisor.\n"
        "For professional consulting, contact joshua@mtm.now."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = MTM_LIGHT
    run.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
